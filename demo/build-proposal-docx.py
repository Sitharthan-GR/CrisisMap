#!/usr/bin/env python3
"""Generate the revised CrisisMap proposal document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/Users/sidhu/Documents/UNCP/CrisisMap/demo/CrisisMap-Proposal.docx"


def set_margins(section, top=1.0, bottom=1.0, left=1.1, right=1.1):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def shade_cell(cell, fill="E8F0FE"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    return h


def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, "EEF4FF")
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    r2 = p.add_run(body)
    r2.font.size = Pt(10.5)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    set_margins(section)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CrisisMap")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Georeferenced Damage Assessment for Humanitarian Response")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    doc.add_paragraph()

    for line, bold in [
        ("Project Proposal", False),
        ("Submitted to: UNDP Humanitarian Innovation & Evaluators", False),
        ("Author: Sitharthan “Sid” Ramalingam", True),
        ("University of North Carolina at Pembroke (UNCP)", False),
        ("June 2026", False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.bold = bold

    doc.add_paragraph()
    add_callout(
        doc,
        "Live prototype",
        "Dashboard: https://crisis-map-phi.vercel.app/\n"
        "Report damage: https://crisis-map-phi.vercel.app/report\n"
        "Admin panel: https://crisis-map-phi.vercel.app/admin\n"
        "Narrated demo video and full design document included in the project repository.",
    )

    doc.add_page_break()

    # 1 Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    add_body(
        doc,
        "When a crisis hits, the first hours are lost to fragmented reporting—WhatsApp threads, "
        "paper forms, and spreadsheets that never align on location, severity, or structure type. "
        "By the time coordinators build a shared picture, decisions are already behind the event.",
    )
    add_body(
        doc,
        "CrisisMap is a web-based damage assessment platform built for that reality. Field teams "
        "submit structured, georeferenced reports from any phone—no account required. Coordinators "
        "open crises, triage unlisted submissions, customize forms, and export validated data as "
        "CSV, GeoJSON, or Shapefile for UNDP, OCHA, national agencies, and NGO partners.",
    )
    add_body(
        doc,
        "The MVP is deployed and ready for evaluation: a live prototype on Vercel, a FastAPI "
        "backend on Render, and PostGIS storage via Supabase. It supports offline reporting, "
        "six languages, version history at the same building, and load testing with hundreds "
        "of concurrent reports. This proposal summarizes the problem, approach, and evidence "
        "that CrisisMap is ready for operational pilot use—not a slide-deck concept.",
    )

    # 2 Problem
    add_heading(doc, "2. The Problem", 1)
    add_body(
        doc,
        "Humanitarian damage assessment fails less from lack of will than from lack of structure. "
        "Field data arrives through incompatible channels, without reliable coordinates, without "
        "a common severity scale, and without a way to track how damage at one building changes "
        "across days or aftershocks.",
    )
    add_bullet(doc, "Reporters drop off when tools require accounts, training, or stable connectivity.")
    add_bullet(doc, "Coordinators spend hours reconciling formats before GIS teams can act.")
    add_bullet(doc, "New events appear before admins can formally open a crisis—valid reports get lost.")
    add_bullet(doc, "Partners need export-ready geodata, not screenshots or ad-hoc spreadsheets.")

    # 3 Approach
    add_heading(doc, "3. What CrisisMap Delivers", 1)

    add_heading(doc, "3.1 Field reporting that matches crisis conditions", 2)
    add_body(
        doc,
        "A guided wizard captures severity (minimal / partial / complete), infrastructure type, "
        "hazard category, debris, location, and up to five photos—in under three minutes. "
        "Reporters can use GPS, map tap, or place search. The nearest active crisis is "
        "pre-selected; if none fits, they choose Other and the report routes to an admin "
        "triage queue instead of blocking submission.",
    )

    add_heading(doc, "3.2 Coordinator control without developer bottlenecks", 2)
    add_body(
        doc,
        "Administrators sign in to a password-protected panel to create and close crises, "
        "review KPIs, assign unlisted reports, and export filtered datasets. A drag-and-drop "
        "form builder lets coordinators add event-specific questions—flood depth, evacuation "
        "status, access road condition—without waiting on engineering.",
    )

    add_heading(doc, "3.3 A map that stays readable as volume grows", 2)
    add_body(
        doc,
        "Repeat assessments at the same building (within a 5-metre tolerance) stack as version "
        "history rather than duplicate pins. The map shows the latest validated assessment; "
        "coordinators and analysts can open the full timeline. This is essential when damage "
        "evolves—floods recede, aftershocks worsen structural risk, or access improves over days.",
    )

    add_heading(doc, "3.4 Offline-first by design", 2)
    add_body(
        doc,
        "Reports and photos queue on the device when connectivity fails. When the network "
        "returns, submissions sync automatically—no action required from the reporter. "
        "This targets post-disaster reality: congested cellular networks, rural dead zones, "
        "and teams working from personal phones rather than issued hardware.",
    )

    # 4 Why different
    add_heading(doc, "4. Why This Approach Works", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Design choice"
    hdr[1].text = "Field impact"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
        shade_cell(cell, "D6E4FF")

    rows = [
        ("No account required", "Removes the largest drop-off point during emergencies"),
        ("Structured wizard + UNDRR-aligned taxonomy", "Comparable statistics across regions and agencies"),
        ("Unlisted report queue", "Nothing is lost when a new event precedes formal crisis setup"),
        ("CSV / GeoJSON / Shapefile export", "Feeds Excel, QGIS, ArcGIS, and coordination dashboards directly"),
        ("Six languages (EN full; AR, FR, ES, ZH, RU partial)", "Serves multinational teams without manual setup"),
        ("PWA + IndexedDB offline sync", "Works on any modern browser—no app store dependency"),
    ]
    for a, b in rows:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
    doc.add_paragraph()

    # 5 Technical foundation
    add_heading(doc, "5. Technical Foundation", 1)
    add_body(
        doc,
        "CrisisMap uses open standards and managed services chosen for reliability, "
        "cost predictability, and alignment with humanitarian IT practice:",
    )
    stack = doc.add_table(rows=1, cols=3)
    stack.style = "Table Grid"
    h = stack.rows[0].cells
    h[0].text = "Layer"
    h[1].text = "Technology"
    h[2].text = "Role"
    for cell in h:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
        shade_cell(cell, "D6E4FF")

    for layer, tech, role in [
        ("Frontend", "React + Vite (PWA)", "Mobile-first UI; installable; offline-capable"),
        ("Backend", "FastAPI on Render.com", "Versioned REST API; managed TLS and deployment"),
        ("Database", "PostgreSQL + PostGIS (Supabase)", "Geospatial queries; location deduplication"),
        ("Storage", "Supabase Storage", "Private photos via time-limited signed URLs"),
        ("Hosting", "Vercel CDN", "Global HTTPS delivery for the dashboard"),
        ("Mapping", "OpenStreetMap + building footprints", "Open data; no licensing fees"),
    ]:
        row = stack.add_row().cells
        row[0].text = layer
        row[1].text = tech
        row[2].text = role
    doc.add_paragraph()

    add_heading(doc, "5.1 Performance and scale", 2)
    add_body(
        doc,
        "The platform has been benchmarked with the included load-testing scripts "
        "(seed_test_reports.py, benchmark_api.py). Map rendering uses optimized stored "
        "procedures and PostGIS indexing so dashboard load remains a single database call "
        "regardless of pin count. Volume tests with 500+ seeded reports confirm map readability, "
        "export integrity, and offline sync behavior. Detailed methodology and screenshots "
        "are documented in docs/UNDP-DESIGN-DOCUMENT.md.",
    )

    add_heading(doc, "5.2 Security", 2)
    add_bullet(doc, "Public reporting: anonymous by default; optional reporter name only.")
    add_bullet(doc, "Admin access: HMAC-signed session token (24 h); sessionStorage only.")
    add_bullet(doc, "Database: service role key server-side; frontend never queries directly.")
    add_bullet(doc, "Photos: signed URLs with short expiry; not publicly listable.")
    add_bullet(doc, "Transport: TLS on Render, Vercel, and Supabase.")

    # 6 Evaluation
    add_heading(doc, "6. Evaluation & Next Steps", 1)
    add_body(
        doc,
        "UNDP evaluators can review CrisisMap immediately through the live deployment. "
        "The repository includes a narrated two-minute demo video, UI screenshots, seed data "
        "for six crisis scenarios (earthquake, flood, wildfire, explosion, conflict, and "
        "closed-crisis export), and a step-by-step evaluator guide. Admin credentials are "
        "available upon request.",
    )
    add_bullet(doc, "Live dashboard: https://crisis-map-phi.vercel.app/")
    add_bullet(doc, "Submit a report: https://crisis-map-phi.vercel.app/report")
    add_bullet(doc, "Design document: docs/UNDP-DESIGN-DOCUMENT.md (in repository)")
    add_bullet(doc, "Demo video: demo/output/CrisisMap_demo_2min_updated_history_createform_with_voiceover.mp4")

    add_heading(doc, "6.1 Roadmap (summary)", 2)
    roadmap = doc.add_table(rows=1, cols=3)
    roadmap.style = "Table Grid"
    rh = roadmap.rows[0].cells
    rh[0].text = "Phase"
    rh[1].text = "Focus"
    rh[2].text = "Target"
    for cell in rh:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
        shade_cell(cell, "D6E4FF")

    for phase, focus, target in [
        ("Production (current)", "Reporting, admin, offline sync, exports, version history, form builder", "June 2026"),
        ("Phase 2", "Per-user admin accounts, MFA, Row-Level Security, full admin i18n", "Q3 2026"),
        ("Phase 3", "OCHA HDX publishing, SMS submission for feature phones", "Q4 2026"),
        ("Phase 4", "Multi-agency data isolation, audit logging, retention policies", "Q1 2027"),
    ]:
        row = roadmap.add_row().cells
        row[0].text = phase
        row[1].text = focus
        row[2].text = target
    doc.add_paragraph()

    # Conclusion
    add_heading(doc, "7. Conclusion", 1)
    add_body(
        doc,
        "CrisisMap closes a practical gap: fast, structured, georeferenced damage reporting "
        "that works on any phone, in multiple languages, with or without connectivity, and "
        "without requiring institutional affiliation to participate. It gives coordinators a "
        "single validated dataset and gives partners export formats they already use.",
    )
    add_body(
        doc,
        "I welcome the opportunity to walk UNDP evaluators through the live prototype, discuss "
        "deployment scenarios, and incorporate field feedback into the next phase of development.",
    )

    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(18)
    r = sig.add_run("Respectfully submitted,\n\n")
    r.font.size = Pt(11)
    r = sig.add_run("Sitharthan “Sid” Ramalingam\n")
    r.bold = True
    r.font.size = Pt(12)
    r = sig.add_run("CrisisMap — University of North Carolina at Pembroke\n")
    r.font.size = Pt(11)
    r = sig.add_run("crisis-map-phi.vercel.app  |  June 2026")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
